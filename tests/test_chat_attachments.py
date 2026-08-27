# -*- coding: utf-8 -*-
"""Offline coverage for ephemeral chat attachment upload and context injection."""

import io
import os
from datetime import datetime, timedelta, timezone

from docx import Document

import config
import main
from layers import attachments, auth, converter, execution, files_store
from layers.converter import ConversionResult, ConversionStatus


def _docx_bytes(text):
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("Attachment", level=1)
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


def _success_state(mode, extra_context):
    return {
        "response": "context received",
        "citations": [],
        "error": "",
        "layer_trace": [],
        "decision_reasoning": "attachment context" if mode == "expert" else None,
        "captured_context": extra_context,
    }


def test_directly_supported_attachment_is_parsed_and_temp_file_removed(
    client,
    auth_headers,
    tmp_path,
    monkeypatch,
):
    monkeypatch.setattr(config, "BASE_DIR", str(tmp_path))
    headers, user = auth_headers("customer")
    uploaded_bytes = _docx_bytes("Unique attachment fact 4821")
    response = client.post(
        "/chat/attachments",
        headers=headers,
        data={"session_id": "attachment-direct"},
        files={"file": ("notes.docx", uploaded_bytes)},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["success"] is True
    assert payload["original_filename"] == "notes.docx"
    record = attachments.get_attachment("attachment-direct", payload["attachment_id"])
    assert record is not None
    assert record.file_id
    assert "Unique attachment fact 4821" in record.text
    stored = files_store.list_files(user["user_id"])
    assert len(stored) == 1
    assert stored[0].source_type == "attachment"
    assert stored[0].session_id == "attachment-direct"
    assert stored[0].file_id == record.file_id
    assert open(files_store.get_file_path(stored[0]), "rb").read() == uploaded_bytes
    assert not (tmp_path / "data" / "tmp_uploads").exists() or not list(
        (tmp_path / "data" / "tmp_uploads").iterdir()
    )
    attachments.clear_session("attachment-direct")


def test_convertible_attachment_is_converted_parsed_and_cleaned(
    client,
    auth_headers,
    tmp_path,
    monkeypatch,
):
    monkeypatch.setattr(config, "BASE_DIR", str(tmp_path))
    headers, user = auth_headers("customer")
    observed = {}

    def fake_convert(source_path, target_format):
        observed["source_path"] = source_path
        output_dir = os.path.join(os.path.dirname(source_path), "conversion_attachment")
        os.makedirs(output_dir)
        output_path = os.path.join(output_dir, "converted.docx")
        with open(output_path, "wb") as output:
            output.write(_docx_bytes("Converted legacy attachment fact"))
        return ConversionResult(
            success=True,
            status=ConversionStatus.SUCCESS,
            output_path=output_path,
            converted_from_format="doc",
            converted_to_format=target_format,
        )

    monkeypatch.setattr(converter, "convert_file", fake_convert)
    response = client.post(
        "/chat/attachments",
        headers=headers,
        data={"session_id": "attachment-converted"},
        files={
            "file": (
                "legacy.doc",
                b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1legacy-content",
            )
        },
    )

    assert response.status_code == 200
    record = attachments.get_attachment(
        "attachment-converted",
        response.json()["attachment_id"],
    )
    assert record is not None
    assert "Converted legacy attachment fact" in record.text
    stored = files_store.list_files(user["user_id"])
    assert stored[0].format == "doc"
    assert stored[0].original_filename == "legacy.doc"
    assert open(files_store.get_file_path(stored[0]), "rb").read().startswith(
        b"\xd0\xcf\x11\xe0"
    )
    assert not os.path.exists(observed["source_path"])
    assert not os.path.exists(os.path.dirname(observed["source_path"])) or not list(
        os.scandir(os.path.dirname(observed["source_path"]))
    )
    attachments.clear_session("attachment-converted")


def test_attachment_rejects_unsupported_format(client, auth_headers):
    headers, _ = auth_headers("customer")
    response = client.post(
        "/chat/attachments",
        headers=headers,
        data={"session_id": "attachment-unsupported"},
        files={"file": ("sheet.csv", b"a,b\n1,2")},
    )

    assert response.status_code == 400
    assert response.json()["error_type"] == "unsupported_format"


def test_attachment_rejects_extracted_text_over_limit(
    client,
    auth_headers,
    monkeypatch,
):
    monkeypatch.setattr(config, "CHAT_ATTACHMENT_MAX_CHARS", 5)
    headers, _ = auth_headers("customer")
    response = client.post(
        "/chat/attachments",
        headers=headers,
        data={"session_id": "attachment-large-text"},
        files={"file": ("notes.txt", b"123456")},
    )

    assert response.status_code == 422
    assert response.json()["error_type"] == "content_too_large"


def test_cross_session_attachment_access_is_rejected(
    client,
    auth_headers,
    monkeypatch,
):
    owner_headers, owner = auth_headers("customer")
    other_headers, other = auth_headers("customer")
    auth.bind_session("attachment-owner", owner["user_id"])
    auth.bind_session("attachment-other", other["user_id"])
    record = attachments.save_attachment(
        "attachment-owner",
        "private attachment text",
        "private.txt",
    )
    monkeypatch.setattr(
        main.planning,
        "run_graph_state",
        lambda session_id, message, mode, extra_context=None: _success_state(
            mode, extra_context
        ),
    )

    response = client.post(
        "/chat",
        headers=other_headers,
        json={
            "session_id": "attachment-other",
            "message": "summarize",
            "mode": "fast",
            "attachment_ids": [record.attachment_id],
        },
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "附件已过期或不存在，请重新上传"
    attachments.clear_session("attachment-owner")
    attachments.clear_session("attachment-other")


def test_expired_attachment_is_removed_lazily(monkeypatch):
    record = attachments.save_attachment("attachment-expired", "old text", "old.txt")
    with attachments._attachment_lock:
        attachments._attachments["attachment-expired"][record.attachment_id].created_at = (
            datetime.now(timezone.utc) - timedelta(minutes=31)
        )
    monkeypatch.setattr(config, "CHAT_ATTACHMENT_TTL_MINUTES", 30)

    assert attachments.get_attachment("attachment-expired", record.attachment_id) is None
    with attachments._attachment_lock:
        assert "attachment-expired" not in attachments._attachments


def test_document_execution_uses_supplied_attachment_context(monkeypatch):
    monkeypatch.setattr(execution.auth, "get_verified_doc_ids", lambda: [])
    monkeypatch.setattr(execution.memory, "search_documents", lambda *args, **kwargs: [])
    observed = {}

    def fake_llm_chat(message, **kwargs):
        observed["message"] = message
        observed["system_prompt"] = kwargs.get("system_prompt", "")
        return "SKY-739 is in the attachment"

    monkeypatch.setattr(execution, "_llm_chat", fake_llm_chat)
    result = execution._search_documents(
        "summarize",
        tier="expert",
        context=["attachment marker SKY-739"],
    )

    assert result.status == "success"
    assert "SKY-739" in result.data
    assert "attachment marker SKY-739" in observed["system_prompt"]
    assert result.metadata["supplied_context_answer"] is True


def test_fast_and_expert_receive_attachment_context(
    client,
    auth_headers,
    monkeypatch,
):
    headers, user = auth_headers("customer")
    captured = []

    def fake_run(
        session_id, message, mode, extra_context=None, owner_user_id="",
        attachment_ids=None, tool_event_sink=None,
    ):
        captured.append((mode, extra_context))
        return _success_state(mode, extra_context)

    monkeypatch.setattr(main.planning, "run_graph_state", fake_run)
    monkeypatch.setattr(main.memory, "save_message", lambda *args: None)
    monkeypatch.setattr(main.memory, "maybe_save_to_vector", lambda *args: None)
    for mode in ("fast", "expert"):
        session_id = "attachment-%s" % mode
        auth.bind_session(session_id, user["user_id"])
        record = attachments.save_attachment(
            session_id,
            "Mode context marker %s" % mode,
            "%s.txt" % mode,
        )
        response = client.post(
            "/chat",
            headers=headers,
            json={
                "session_id": session_id,
                "message": "summarize",
                "mode": mode,
                "attachment_ids": [record.attachment_id],
            },
        )
        assert response.status_code == 200
        attachments.clear_session(session_id)

    assert [item[0] for item in captured] == ["fast", "expert"]
    assert "Mode context marker fast" in captured[0][1][0]
    assert "Mode context marker expert" in captured[1][1][0]


def test_fast_stream_receives_attachment_context(
    client,
    auth_headers,
    monkeypatch,
):
    headers, user = auth_headers("customer")
    session_id = "attachment-fast-stream"
    auth.bind_session(session_id, user["user_id"])
    record = attachments.save_attachment(
        session_id,
        "Stream attachment marker 8642",
        "stream.txt",
    )
    captured = {}

    def fake_run(
        session_id, message, mode, extra_context=None, owner_user_id="",
        attachment_ids=None, tool_event_sink=None,
    ):
        captured["context"] = extra_context
        return _success_state(mode, extra_context)

    monkeypatch.setattr(main.planning, "run_graph_state", fake_run)
    monkeypatch.setattr(main.memory, "save_message", lambda *args: None)
    monkeypatch.setattr(main.memory, "maybe_save_to_vector", lambda *args: None)
    response = client.post(
        "/chat/stream",
        headers=headers,
        json={
            "session_id": session_id,
            "message": "summarize",
            "mode": "fast",
            "attachment_ids": [record.attachment_id],
        },
    )

    assert response.status_code == 200
    assert '"chunk": "[DONE]"' in response.text
    assert "Stream attachment marker 8642" in captured["context"][0]
    attachments.clear_session(session_id)
