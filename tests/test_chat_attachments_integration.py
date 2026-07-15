# -*- coding: utf-8 -*-
"""Real DeepSeek chat attachment coverage; excluded from CI."""

import io

import pytest
from docx import Document

from layers import attachments, memory


pytestmark = [pytest.mark.integration, pytest.mark.slow]


def _real_docx_bytes():
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("苍穹项目内部备忘", level=1)
    document.add_paragraph("项目代号为SKY-739。计划发布日期为9月18日。负责人是林岚。")
    document.add_paragraph("发布前必须完成权限审计和数据备份演练。")
    document.save(buffer)
    return buffer.getvalue()


def test_real_fast_and_expert_read_uploaded_docx(client, auth_headers):
    headers, _ = auth_headers("customer")
    session_ids = ["attachment-real-fast", "attachment-real-expert"]
    try:
        for mode, session_id in zip(("fast", "expert"), session_ids):
            uploaded = client.post(
                "/chat/attachments",
                headers=headers,
                data={"session_id": session_id},
                files={"file": ("project_memo.docx", _real_docx_bytes())},
            )
            assert uploaded.status_code == 200, uploaded.text
            attachment_id = uploaded.json()["attachment_id"]

            response = client.post(
                "/chat",
                headers=headers,
                json={
                    "session_id": session_id,
                    "message": "请总结这个文件，并明确写出项目代号、发布日期和负责人。",
                    "mode": mode,
                    "attachment_ids": [attachment_id],
                },
                timeout=180,
            )
            assert response.status_code == 200, response.text
            payload = response.json()
            assert payload["status"] == "success", payload
            assert "SKY-739" in payload["data"]
            assert "9月18" in payload["data"]
            assert "林岚" in payload["data"]
    finally:
        for session_id in session_ids:
            attachments.clear_session(session_id)
            memory.clear_session(session_id)
