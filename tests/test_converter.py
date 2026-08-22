# -*- coding: utf-8 -*-
"""Offline tests for serialized LibreOffice conversion."""

import os
import subprocess
from types import SimpleNamespace

import config
from docx import Document
from layers import converter, execution, planning


def _source_file(tmp_path):
    source = tmp_path / "input.doc"
    source.write_bytes(b"office source")
    return source


def test_convert_file_success(tmp_path, monkeypatch):
    source = _source_file(tmp_path)
    monkeypatch.setattr(config, "LIBREOFFICE_PATH", str(tmp_path / "soffice.exe"))
    (tmp_path / "soffice.exe").write_bytes(b"stub")

    def run(command, **kwargs):
        output_dir = command[command.index("--outdir") + 1]
        output_path = os.path.join(output_dir, "input.docx")
        document = Document()
        document.add_paragraph("converted")
        document.save(output_path)
        return SimpleNamespace(returncode=0, stdout=b"", stderr=b"")

    monkeypatch.setattr(converter.subprocess, "run", run)
    result = converter.convert_file(str(source), "docx")

    assert result.success is True
    assert result.status == converter.ConversionStatus.SUCCESS
    assert result.converted_from_format == "doc"
    assert result.converted_to_format == "docx"
    assert result.error_type == ""
    assert result.output_path and os.path.isfile(result.output_path)
    converter.cleanup_conversion_output(result.output_path)
    assert not [path for path in tmp_path.iterdir() if path.name.startswith("conversion_")]


def test_convert_file_failure_cleans_partial_output(tmp_path, monkeypatch):
    source = _source_file(tmp_path)
    monkeypatch.setattr(config, "LIBREOFFICE_PATH", str(tmp_path / "soffice.exe"))
    (tmp_path / "soffice.exe").write_bytes(b"stub")

    def run(command, **kwargs):
        output_dir = command[command.index("--outdir") + 1]
        with open(os.path.join(output_dir, "partial.docx"), "wb") as output:
            output.write(b"partial")
        return SimpleNamespace(returncode=1, stdout=b"", stderr=b"failed")

    monkeypatch.setattr(converter.subprocess, "run", run)
    result = converter.convert_file(str(source), "docx")

    assert result.success is False
    assert result.status == converter.ConversionStatus.FAILED
    assert result.error_type == "process_failed"
    assert not [path for path in tmp_path.iterdir() if path.name.startswith("conversion_")]


def test_convert_file_timeout_cleans_partial_output(tmp_path, monkeypatch):
    source = _source_file(tmp_path)
    monkeypatch.setattr(config, "LIBREOFFICE_PATH", str(tmp_path / "soffice.exe"))
    (tmp_path / "soffice.exe").write_bytes(b"stub")

    def run(command, **kwargs):
        output_dir = command[command.index("--outdir") + 1]
        with open(os.path.join(output_dir, "partial.docx"), "wb") as output:
            output.write(b"partial")
        raise subprocess.TimeoutExpired(command, kwargs["timeout"])

    monkeypatch.setattr(converter.subprocess, "run", run)
    result = converter.convert_file(str(source), "docx")

    assert result.success is False
    assert result.status == converter.ConversionStatus.TIMEOUT
    assert result.error_type == "timeout"
    assert not [path for path in tmp_path.iterdir() if path.name.startswith("conversion_")]


def test_convert_document_is_registered_for_expert_only():
    assert execution.TOOL_REGISTRY["convert_document"] == "_convert_document"
    exposed_tools = {
        function["function"]["name"]
        for function in planning.INTENT_TOOLS
        if function.get("function")
    }
    fast_tools = {
        function["function"]["name"]
        for function in planning.FAST_TOOLS
        if function.get("function")
    }
    assert "convert_document" in exposed_tools
    assert "convert_document" not in fast_tools


def test_pdf_reconstruction_does_not_share_libreoffice_lock():
    assert converter._pdf_conversion_lock is not converter._conversion_lock


def test_libreoffice_wrapper_matches_legacy_success_result(tmp_path, monkeypatch):
    source = _source_file(tmp_path)
    monkeypatch.setattr(config, "LIBREOFFICE_PATH", str(tmp_path / "soffice.exe"))
    (tmp_path / "soffice.exe").write_bytes(b"stub")

    def run(command, **kwargs):
        output_dir = command[command.index("--outdir") + 1]
        output_path = os.path.join(output_dir, "input.docx")
        document = Document()
        document.add_paragraph("中文转换对比")
        document.save(output_path)
        return SimpleNamespace(returncode=0, stdout=b"", stderr=b"")

    monkeypatch.setattr(converter.subprocess, "run", run)
    legacy = converter._convert_file_impl(str(source), "docx")
    wrapped = converter.convert_file(str(source), "docx")

    assert wrapped.model_dump(exclude={"output_path"}) == legacy.model_dump(
        exclude={"output_path"}
    )
    assert Document(wrapped.output_path).paragraphs[0].text == "中文转换对比"
    converter.cleanup_conversion_output(legacy.output_path or "")
    converter.cleanup_conversion_output(wrapped.output_path or "")


def test_libreoffice_wrapper_matches_legacy_failure_result(tmp_path):
    missing = str(tmp_path / "missing.doc")

    legacy = converter._convert_file_impl(missing, "docx")
    wrapped = converter.convert_file(missing, "docx")

    assert wrapped == legacy
